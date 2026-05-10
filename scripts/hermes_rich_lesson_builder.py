#!/usr/bin/env python3
import argparse, json, os, re, subprocess, zipfile
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
VERIFY=Path('/root/.hermes/scripts/hermes_verify.py'); MEDIA_VERIFY=Path('/root/.hermes/scripts/hermes_media_verify.py'); DELIVER=Path('/root/.hermes/scripts/hermes_telegram_deliver.py')
PAL={'blue':(80,150,255),'sky':(175,220,255),'green':(90,185,125),'yellow':(255,215,90),'red':(235,85,105),'pink':(255,185,215),'purple':(160,120,255),'brown':(120,80,45),'skin':(245,190,150),'dark':(35,45,65),'white':(255,255,255)}
def run(cmd,timeout=180): return subprocess.run(cmd,shell=True,text=True,capture_output=True,timeout=timeout)
def font(size,bold=False):
    for p in ['/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf' if bold else '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf','/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf' if bold else '/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf']:
        if Path(p).exists(): return ImageFont.truetype(p,size=size)
    return ImageFont.load_default()
def slug(text): return re.sub(r'[^a-z0-9]+','_', (text or 'item').lower()).strip('_')[:60] or 'item'
def canvas(title,bg=(245,250,255)):
    img=Image.new('RGB',(900,560),bg); d=ImageDraw.Draw(img); d.rounded_rectangle((24,24,876,536),radius=40,fill=(255,255,255),outline=(180,205,255),width=7); d.text((44,36),title[:24],font=font(48,True),fill=PAL['dark']); return img,d
def person(d,x,y,shirt):
    d.ellipse((x+55,y+10,x+145,y+100),fill=PAL['skin'],outline=PAL['dark'],width=4); d.ellipse((x+78,y+45,x+90,y+57),fill=PAL['dark']); d.ellipse((x+110,y+45,x+122,y+57),fill=PAL['dark']); d.arc((x+82,y+58,x+122,y+85),10,170,fill=PAL['red'],width=4); d.rounded_rectangle((x+45,y+110,x+155,y+240),radius=30,fill=shirt,outline=PAL['dark'],width=4)
def img_home(path,title):
    img,d=canvas(title); d.polygon([(210,300),(450,135),(700,300)],fill=PAL['red'],outline=PAL['dark']); d.rounded_rectangle((260,300,650,505),radius=18,fill=(255,240,200),outline=PAL['dark'],width=5); d.rectangle((420,390,500,505),fill=PAL['brown'],outline=PAL['dark'],width=4); d.rectangle((310,345,390,420),fill=PAL['sky'],outline=PAL['dark'],width=4); d.rectangle((530,345,610,420),fill=PAL['sky'],outline=PAL['dark'],width=4); d.text((345,215),'HOME',font=font(54,True),fill=PAL['dark']); img.save(path)
def img_family(path,title):
    img,d=canvas(title,(255,246,250)); person(d,175,210,PAL['blue']); person(d,350,170,PAL['pink']); person(d,525,220,PAL['green']); d.text((315,465),'Family',font=font(54,True),fill=PAL['purple']); img.save(path)
def img_book(path,title):
    img,d=canvas(title,(250,245,255)); d.polygon([(210,220),(450,160),(450,455),(210,510)],fill=PAL['blue'],outline=PAL['dark']); d.polygon([(450,160),(700,220),(700,510),(450,455)],fill=PAL['green'],outline=PAL['dark']); d.line((450,160,450,455),fill=PAL['dark'],width=5); d.text((310,92),'Read',font=font(58,True),fill=PAL['purple']); img.save(path)
def img_door(path,title):
    img,d=canvas(title,(245,255,250)); d.rectangle((240,190,410,500),fill=PAL['brown'],outline=PAL['dark'],width=5); d.ellipse((370,340,390,360),fill=PAL['yellow'],outline=PAL['dark']); d.rectangle((520,210,720,390),fill=PAL['sky'],outline=PAL['dark'],width=5); d.line((620,210,620,390),fill=PAL['dark'],width=4); d.line((520,300,720,300),fill=PAL['dark'],width=4); d.text((260,95),'Door + Window',font=font(44,True),fill=PAL['green']); img.save(path)
def img_generic(path,title):
    img,d=canvas(title,(250,250,255)); d.ellipse((300,180,600,480),fill=PAL['yellow'],outline=PAL['dark'],width=6); d.text((330,295),title[:12],font=font(48,True),fill=PAL['dark']); img.save(path)
def gen_img(path,title,topic):
    tier=os.environ.get('HERMES_VISUAL_QUALITY_TIER','local_cartoon_scene')
    cmd=['/usr/bin/python3','/root/.hermes/scripts/hermes_visual_asset_router.py','--label',str(title),'--topic',str(topic),'--output',str(path),'--quality-tier',tier]
    proc=subprocess.run(cmd,text=True,capture_output=True,timeout=180)
    if proc.returncode:
        raise RuntimeError('VISUAL_ASSET_ROUTER_FAILED\n'+proc.stdout+proc.stderr)
    verify=subprocess.run(['/usr/bin/python3','/root/.hermes/scripts/hermes_visual_quality_verify.py','--image',str(path),'--min-tier',tier],text=True,capture_output=True,timeout=120)
    if verify.returncode:
        raise RuntimeError('VISUAL_QUALITY_VERIFY_FAILED\n'+verify.stdout+verify.stderr)
def sections(topic,age,path=''):
    if path:
        data=json.loads(Path(path).read_text()); return [(x.get('title',f'Section {i+1}'),x.get('sentence','')) for i,x in enumerate(data)]
    if 'home' in topic.lower(): return [('My Home','This is my home.'),('My House','My house is safe and cozy.'),('Door','This is my door.'),('Window','This is my window.'),('Mommy','This is my mommy.'),('Daddy','This is my daddy.'),('Book','This is my book.'),('Review','Point to the door. Point to the window. Say: This is my home.')]
    return [(topic.title(),f'This lesson is about {topic}.'),('Learn','We can learn new words.'),('Say','We can say simple sentences.'),('Review','Let us review together.')]
def build_docx(out,title,topic,age,secs,imgdir):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    out=Path(out); out.parent.mkdir(parents=True,exist_ok=True); imgdir=Path(imgdir); imgdir.mkdir(parents=True,exist_ok=True); doc=Document(); doc.styles['Normal'].font.name='Arial'; doc.styles['Normal'].font.size=Pt(12); h=doc.add_heading(title,0); h.alignment=WD_ALIGN_PARAGRAPH.CENTER; p=doc.add_paragraph(f'Age: {age}'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; imgs=[]
    for i,(st,sent) in enumerate(secs,1):
        ip=imgdir/f'{i:02d}_{slug(st)}.png'; gen_img(ip,st,topic); imgs.append(ip); doc.add_heading(st,1); pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.add_run().add_picture(str(ip),width=Inches(4.6)); sp=doc.add_paragraph(sent); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.save(out); assert zipfile.is_zipfile(out); return imgs
def build_pdf(out, title, topic, age, secs, imgdir):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, PageBreak, KeepTogether
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    out = Path(out)
    out.parent.mkdir(parents=True, exist_ok=True)
    imgdir = Path(imgdir)
    imgdir.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(
        str(out),
        pagesize=A4,
        rightMargin=0.48 * inch,
        leftMargin=0.48 * inch,
        topMargin=0.42 * inch,
        bottomMargin=0.42 * inch,
    )
    styles = getSampleStyleSheet()
    ts = ParagraphStyle('LessonTitle', parent=styles['Title'], fontSize=22, leading=26, alignment=1, textColor=colors.HexColor('#0f172a'), spaceAfter=4)
    age_style = ParagraphStyle('AgeStyle', parent=styles['BodyText'], fontSize=10, leading=12, alignment=1, textColor=colors.HexColor('#334155'), spaceAfter=8)
    hs = ParagraphStyle('LessonHeading', parent=styles['Heading2'], fontSize=15, leading=18, alignment=1, textColor=colors.HexColor('#be185d'), spaceAfter=4)
    bs = ParagraphStyle('LessonBody', parent=styles['BodyText'], fontSize=11, leading=14, alignment=1, textColor=colors.HexColor('#1f2937'), spaceAfter=2)

    story = [Paragraph(title, ts), Paragraph(f'Age: {age}', age_style), Spacer(1, 0.08 * inch)]
    imgs = []
    cards = []
    for i, (st, sent) in enumerate(secs, 1):
        ip = imgdir / f'{i:02d}_{slug(st)}.png'
        gen_img(ip, st, topic)
        imgs.append(ip)
        cards.append(KeepTogether([
            Paragraph(st, hs),
            RLImage(str(ip), width=4.25 * inch, height=2.8 * inch),
            Spacer(1, 0.05 * inch),
            Paragraph(sent, bs),
        ]))

    for idx, card in enumerate(cards, 1):
        story.append(card)
        story.append(Spacer(1, 0.16 * inch))
        if idx % 2 == 0 and idx != len(cards):
            story.append(PageBreak())

    doc.build(story)
    if not out.read_bytes().startswith(b'%PDF'):
        raise RuntimeError('PDF_INVALID')
    return imgs

def verify(out,fmt,min_media):
    vf=run(f"{VERIFY} file --path {json.dumps(str(out))}");
    if vf.returncode: raise RuntimeError(vf.stdout+vf.stderr)
    mv=run(f"{MEDIA_VERIFY} --file {json.dumps(str(out))} --type {fmt} --min-media {int(min_media)}")
    if mv.returncode: raise RuntimeError(mv.stdout+mv.stderr)
def deliver(out,fmt):
    proc=run(f"{DELIVER} --file {json.dumps(str(out))} --caption {json.dumps('Your Majesty, here is your verified rich '+fmt.upper()+' lesson with pictures: '+Path(out).name)} --preview-link no",timeout=180)
    if proc.returncode: raise RuntimeError(proc.stdout+proc.stderr)
def main():
    ap=argparse.ArgumentParser(description='Create rich PDF/DOCX lessons with local pictures and media verification.'); ap.add_argument('--format',required=True,choices=['pdf','docx']); ap.add_argument('--output',required=True); ap.add_argument('--title',required=True); ap.add_argument('--topic',default=''); ap.add_argument('--age',default=''); ap.add_argument('--sections-json',default=''); ap.add_argument('--min-media',type=int,default=1); ap.add_argument('--quality-tier',default='local_cartoon_scene'); ap.add_argument('--deliver-telegram',action='store_true'); ap.add_argument('--no-deliver-telegram',action='store_true'); a=ap.parse_args(); topic=a.topic or a.title; age=a.age or 'kids'; os.environ['HERMES_VISUAL_QUALITY_TIER']=a.quality_tier; secs=sections(topic,age,a.sections_json); imgdir=Path('/root/.hermes/file_outputs/rich_lesson_images')/Path(a.output).stem; imgs=build_pdf(a.output,a.title,topic,age,secs,imgdir) if a.format=='pdf' else build_docx(a.output,a.title,topic,age,secs,imgdir); min_media=a.min_media or len(imgs); verify(a.output,a.format,min_media); status='SKIPPED'
    if not a.no_deliver_telegram or a.deliver_telegram: deliver(a.output,a.format); status='PASSED'
    print('RICH_LESSON_BUILD=PASSED'); print('FORMAT='+a.format); print('OUTPUT_PATH='+str(a.output)); print('SECTION_COUNT='+str(len(secs))); print('GENERATED_IMAGE_COUNT='+str(len(imgs))); print('IMAGES_DIR='+str(imgdir)); print('MIN_MEDIA_REQUIRED='+str(min_media)); print('VERIFICATION=PASSED'); print('MEDIA_VERIFICATION=PASSED'); print('TELEGRAM_DELIVERY='+status)
if __name__=='__main__': main()
